"""
多格式文件解析服务
支持 .txt / .md / .docx / .pdf
"""
import os
from typing import Optional
from app.services.llm_adapter import get_llm
from app.db import crud


class MaterialParser:
    """文件解析与语料提取"""
    
    SUPPORTED_TYPES = {".txt", ".md", ".docx", ".pdf"}
    
    def parse_file(self, filename: str, content_bytes: bytes) -> dict:
        """
        解析文件内容
        返回: {"raw_text": str, "sections": list, "word_count": int}
        """
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == ".txt":
            return self._parse_txt(content_bytes)
        elif ext == ".md":
            return self._parse_markdown(content_bytes)
        elif ext == ".docx":
            return self._parse_docx(content_bytes)
        elif ext == ".pdf":
            return self._parse_pdf(content_bytes)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _parse_txt(self, content_bytes: bytes) -> dict:
        """解析纯文本"""
        text = content_bytes.decode("utf-8", errors="ignore")
        return {
            "raw_text": text,
            "sections": [{"title": "Full Text", "content": text}],
            "word_count": len(text.split())
        }
    
    def _parse_markdown(self, content_bytes: bytes) -> dict:
        """解析 Markdown（提取标题分段）"""
        text = content_bytes.decode("utf-8", errors="ignore")
        sections = []
        current_title = "Introduction"
        current_content = []
        
        for line in text.split("\n"):
            if line.startswith("#"):
                if current_content:
                    sections.append({"title": current_title, "content": "\n".join(current_content)})
                current_title = line.lstrip("#").strip()
                current_content = []
            else:
                current_content.append(line)
        
        if current_content:
            sections.append({"title": current_title, "content": "\n".join(current_content)})
        
        return {
            "raw_text": text,
            "sections": sections,
            "word_count": len(text.split())
        }
    
    def _parse_docx(self, content_bytes: bytes) -> dict:
        """解析 Word 文档"""
        import io
        try:
            from docx import Document
            doc = Document(io.BytesIO(content_bytes))
            
            paragraphs = []
            sections = []
            current_title = "Document"
            current_content = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                paragraphs.append(text)
                
                # 检测标题段落
                if para.style and "heading" in para.style.name.lower():
                    if current_content:
                        sections.append({"title": current_title, "content": "\n".join(current_content)})
                    current_title = text
                    current_content = []
                else:
                    current_content.append(text)
            
            if current_content:
                sections.append({"title": current_title, "content": "\n".join(current_content)})
            
            # 也提取表格内容
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    sections.append({"title": "Table", "content": "\n".join(table_text)})
                    paragraphs.extend(table_text)
            
            full_text = "\n".join(paragraphs)
            return {
                "raw_text": full_text,
                "sections": sections if sections else [{"title": "Document", "content": full_text}],
                "word_count": len(full_text.split())
            }
        except ImportError:
            return {"raw_text": "[DOCX parsing unavailable - python-docx not installed]", "sections": [], "word_count": 0}
        except Exception as e:
            return {"raw_text": f"[DOCX parsing error: {str(e)}]", "sections": [], "word_count": 0}
    
    def _parse_pdf(self, content_bytes: bytes) -> dict:
        """解析 PDF"""
        import io
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content_bytes))
            
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append({"title": f"Page {i+1}", "content": text.strip()})
            
            full_text = "\n\n".join([p["content"] for p in pages])
            return {
                "raw_text": full_text,
                "sections": pages,
                "word_count": len(full_text.split())
            }
        except ImportError:
            return {"raw_text": "[PDF parsing unavailable - PyPDF2 not installed]", "sections": [], "word_count": 0}
        except Exception as e:
            return {"raw_text": f"[PDF parsing error: {str(e)}]", "sections": [], "word_count": 0}
    
    async def analyze_for_corpus(self, parsed_content: dict, target_band: str = "7.0") -> dict:
        """
        使用 LLM 分析解析后的内容，提取可用于语料库的元素
        返回: {"vocabulary": [...], "expressions": [...], "topic_ideas": [...], "anchor_suggestions": [...]}
        """
        llm = get_llm()
        raw_text = parsed_content.get("raw_text", "")
        
        # 截取前3000字符避免token过长
        text_excerpt = raw_text[:3000]
        
        prompt = f"""Analyze the following English learning material and extract useful elements for an IELTS speaking corpus.
Target band: {target_band}

Material content:
---
{text_excerpt}
---

Extract and return as JSON:
{{
  "vocabulary": [
    {{"basic_word": "...", "upgrade": "...", "context": "example sentence", "category": "topic_category"}}
  ],
  "expressions": [
    {{"expression": "...", "meaning": "...", "when_to_use": "..."}}
  ],
  "topic_ideas": [
    {{"topic": "...", "angle": "how this connects to a personal story", "sample_phrase": "..."}}
  ],
  "anchor_suggestions": [
    {{"theme": "...", "story_seed": "a brief story idea based on this material"}}
  ]
}}

Return ONLY valid JSON, no markdown formatting."""

        try:
            response = await llm.chat([
                {"role": "system", "content": "You are an IELTS speaking coach analyzing study materials. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ], temperature=0.3)
            
            # 尝试解析JSON
            import json
            # 清理可能的markdown代码块
            response = response.strip()
            if response.startswith("```"):
                response = response.split("\n", 1)[1] if "\n" in response else response
                response = response.rsplit("```", 1)[0] if "```" in response else response
            
            result = json.loads(response)
            return result
        except Exception as e:
            # Fallback：返回空分析结果
            return {
                "vocabulary": [],
                "expressions": [],
                "topic_ideas": [],
                "anchor_suggestions": [],
                "error": str(e)
            }
    
    async def merge_to_corpus(self, corpus_id: str, analysis_result: dict,
                              selected_items: dict = None) -> dict:
        """
        将分析结果融合到语料库
        selected_items: 用户选择要融合的条目（如果为None则全部融合）
        """
        from app.db.crud import get_corpus, update_corpus
        import json
        
        corpus = await get_corpus(corpus_id)
        if not corpus:
            return {"error": "Corpus not found"}
        
        changes = []
        
        # 融合词汇
        vocab_to_add = selected_items.get("vocabulary", analysis_result.get("vocabulary", [])) if selected_items else analysis_result.get("vocabulary", [])
        if vocab_to_add:
            existing_vocab = corpus.get("vocabulary") or []
            if isinstance(existing_vocab, str):
                existing_vocab = json.loads(existing_vocab)
            existing_vocab.extend(vocab_to_add)
            await update_corpus(corpus_id, {"vocabulary": existing_vocab})
            changes.append(f"Added {len(vocab_to_add)} vocabulary items")
        
        # 融合表达到patterns
        expressions = selected_items.get("expressions", analysis_result.get("expressions", [])) if selected_items else analysis_result.get("expressions", [])
        if expressions:
            existing_patterns = corpus.get("patterns") or []
            if isinstance(existing_patterns, str):
                existing_patterns = json.loads(existing_patterns)
            for expr in expressions:
                existing_patterns.append({
                    "name": expr.get("expression", ""),
                    "formula": expr.get("expression", ""),
                    "example": expr.get("meaning", ""),
                    "when_to_use": expr.get("when_to_use", "")
                })
            await update_corpus(corpus_id, {"patterns": existing_patterns})
            changes.append(f"Added {len(expressions)} expression patterns")
        
        return {"status": "merged", "changes": changes}


def get_material_parser() -> MaterialParser:
    return MaterialParser()
