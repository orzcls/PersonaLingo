"""
Dump docx raw paragraphs + tables to stdout for inspection.
Usage: python backend/tools/inspect_demo_docx.py
"""
from __future__ import annotations
import sys
from pathlib import Path

try:
    from docx import Document  # python-docx
except ImportError:
    print("[inspect] pip install python-docx", file=sys.stderr)
    sys.exit(1)

WORKSPACE_ROOT = Path(__file__).resolve().parents[3]
DEMO_DIR = WORKSPACE_ROOT / "语料库项目demo"

TARGETS = [
    DEMO_DIR / "PPChu冲鸭小分队—地道单词.docx",
    DEMO_DIR / "程同学 课堂精练.docx",
]


def dump(path: Path) -> None:
    print(f"\n================ {path.name} ================")
    if not path.exists():
        print(f"[MISSING] {path}")
        return
    doc = Document(str(path))
    print(f"[paragraphs] {len(doc.paragraphs)}")
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t:
            print(f"P{i:04d}: {t}")
    print(f"\n[tables] {len(doc.tables)}")
    for ti, tbl in enumerate(doc.tables):
        print(f"--- table {ti} ({len(tbl.rows)} rows x {len(tbl.columns)} cols) ---")
        for ri, row in enumerate(tbl.rows):
            cells = [(c.text or "").strip().replace("\n", " | ") for c in row.cells]
            print(f"T{ti}R{ri:03d}: {cells}")


def main() -> int:
    for p in TARGETS:
        dump(p)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
